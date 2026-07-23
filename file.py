"""
Vulnerable Lab - Project Report Generator
Creates a professional DOCX and PDF report
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
from datetime import datetime

class ReportGenerator:
    def __init__(self):
        self.doc = Document()
        self.setup_document()
        
    def setup_document(self):
        """Setup document margins and styles"""
        # Set margins
        sections = self.doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
    
    def add_heading(self, text, level=1):
        """Add a heading with proper formatting"""
        heading = self.doc.add_heading(text, level)
        # Make headings blue for professional look
        for run in heading.runs:
            run.font.color.rgb = RGBColor(0, 51, 102)
        return heading
    
    def add_paragraph(self, text, style='Normal'):
        """Add a paragraph with proper formatting"""
        p = self.doc.add_paragraph(text, style=style)
        return p
    
    def add_table(self, headers, data):
        """Add a formatted table"""
        table = self.doc.add_table(rows=1, cols=len(headers))
        table.style = 'Light Grid Accent 1'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Add headers
        header_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            header_cells[i].text = header
            header_cells[i].paragraphs[0].runs[0].font.bold = True
        
        # Add data
        for row_data in data:
            row = table.add_row()
            for i, cell_data in enumerate(row_data):
                row.cells[i].text = str(cell_data)
        
        return table
    
    def add_screenshot_placeholder(self, figure_num, caption):
        """Add a screenshot placeholder with figure number and caption"""
        # Add placeholder box
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"[SCREENSHOT PLACEHOLDER - FIGURE {figure_num}]")
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = RGBColor(128, 128, 128)
        
        # Add caption
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"Figure {figure_num}: {caption}")
        run.font.size = Pt(10)
        run.font.italic = True
    
    def generate_report(self):
        """Generate the complete report"""
        
        # ============================================
        # TITLE PAGE
        # ============================================
        # Add title with spacing
        self.doc.add_paragraph()
        self.doc.add_paragraph()
        self.doc.add_paragraph()
        
        title = self.doc.add_heading('VULNERABLE LAB', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in title.runs:
            run.font.size = Pt(36)
            run.font.color.rgb = RGBColor(0, 51, 102)
        
        self.doc.add_paragraph()
        subtitle = self.doc.add_paragraph('Custom Vulnerable Web Application')
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in subtitle.runs:
            run.font.size = Pt(18)
            run.font.italic = True
        
        self.doc.add_paragraph()
        self.doc.add_paragraph()
        
        # Project info
        info = [
            ('Project Type:', 'Cybersecurity Education Tool'),
            ('Course:', 'Cybersecurity Project'),
            ('Submission Date:', datetime.now().strftime('%B %d, %Y')),
            ('Author:', 'Your Name'),
            ('Instructor:', 'Your Instructor Name')
        ]
        
        for label, value in info:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(label)
            run.font.bold = True
            run = p.add_run(f' {value}')
        
        # Add page break
        self.doc.add_page_break()
        
        # ============================================
        # TABLE OF CONTENTS
        # ============================================
        self.add_heading('TABLE OF CONTENTS', 1)
        self.doc.add_paragraph('1. Problem Statement ...................... 1')
        self.doc.add_paragraph('2. Project Overview ........................ 2')
        self.doc.add_paragraph('3. Architecture & Technology ................ 3')
        self.doc.add_paragraph('4. Features & Vulnerabilities ............... 4')
        self.doc.add_paragraph('5. Testing & Demonstration ................. 6')
        self.doc.add_paragraph('6. Results & Output ........................ 8')
        self.doc.add_paragraph('7. Limitations ............................ 10')
        self.doc.add_paragraph('8. Future Improvements .................... 11')
        self.doc.add_paragraph('9. Conclusion ............................. 12')
        self.doc.add_paragraph('10. References ............................ 13')
        self.doc.add_page_break()
        
        # ============================================
        # 1. PROBLEM STATEMENT
        # ============================================
        self.add_heading('1. PROBLEM STATEMENT', 1)
        
        self.add_paragraph("""
        In the field of cybersecurity, there is a critical need for safe, controlled environments 
        where security professionals and students can learn about web application vulnerabilities 
        without risking real systems. Traditional security training often lacks hands-on experience 
        with actual vulnerable applications.
        """)
        
        self.add_heading('1.1 Key Challenges', 2)
        self.add_paragraph("""
        • Security professionals need practical experience identifying and exploiting vulnerabilities
        • Traditional training methods are often theoretical and lack real-world application
        • Testing on live systems is unethical and illegal
        • There is a shortage of accessible, educational vulnerable applications
        """)
        
        self.add_heading('1.2 Our Solution', 2)
        self.add_paragraph("""
        We developed Vulnerable Lab - a deliberately insecure web application that contains multiple 
        common vulnerabilities. This tool serves as an educational platform for:
        
        • Learning about OWASP Top 10 vulnerabilities
        • Practicing penetration testing techniques
        • Understanding secure coding practices
        • Developing security testing skills in a safe environment
        """)
        
        self.doc.add_page_break()
        
        # ============================================
        # 2. PROJECT OVERVIEW
        # ============================================
        self.add_heading('2. PROJECT OVERVIEW', 1)
        
        self.add_heading('2.1 What is Vulnerable Lab?', 2)
        self.add_paragraph("""
        Vulnerable Lab is a deliberately insecure web application designed as an educational tool 
        for learning about cybersecurity vulnerabilities. It simulates real-world security flaws 
        so you can practice identifying, exploiting, and fixing them in a safe environment.
        """)
        
        self.add_heading('2.2 Project Purpose', 2)
        self.add_paragraph("""
        • Educational: Teaches common web vulnerabilities
        • Hands-on Practice: Allows safe penetration testing
        • Security Awareness: Demonstrates why secure coding matters
        • Ethical Hacking: Practice in a controlled environment
        """)
        
        self.add_heading('2.3 Target Audience', 2)
        self.add_paragraph("""
        • Cybersecurity students
        • Security professionals
        • Software developers
        • CTF (Capture The Flag) players
        • Security educators
        """)
        
        self.doc.add_page_break()
        
        # ============================================
        # 3. ARCHITECTURE & TECHNOLOGY
        # ============================================
        self.add_heading('3. ARCHITECTURE & TECHNOLOGY', 1)
        
        self.add_heading('3.1 System Architecture', 2)
        
        self.add_paragraph("""
        The application follows a client-server architecture with the following components:
        
        Frontend: HTML5, CSS3, Bootstrap 5 for responsive user interface
        Backend: Python Flask framework for server-side logic
        Database: SQLite with SQLAlchemy ORM for data storage
        Authentication: Flask-Login for session management
        """)
        
        self.add_screenshot_placeholder(1, "System Architecture Diagram")
        
        self.add_heading('3.2 Technology Stack', 2)
        
        tech_data = [
            ('Component', 'Technology', 'Version'),
            ('Backend', 'Python', '3.8+'),
            ('Framework', 'Flask', '2.3.2'),
            ('Database', 'SQLite', '3.x'),
            ('ORM', 'SQLAlchemy', '3.0.5'),
            ('Frontend', 'Bootstrap', '5.1.3'),
            ('Authentication', 'Flask-Login', '0.6.2')
        ]
        self.add_table(['Component', 'Technology', 'Version'], tech_data[1:])
        
        self.add_heading('3.3 Project Structure', 2)
        
        self.add_paragraph("""
        vulnerable_lab/
        ├── app.py              # Main application with all vulnerabilities
        ├── database.py         # Database models
        ├── requirements.txt    # Python dependencies
        ├── templates/          # HTML templates
        │   ├── index.html      # Home page
        │   ├── login.html      # SQL Injection vulnerability
        │   ├── register.html   # Weak Authentication
        │   ├── dashboard.html  # Main dashboard
        │   ├── profile.html    # IDOR vulnerability
        │   ├── search.html     # XSS vulnerability
        │   ├── ping.html       # Command Injection
        │   ├── upload.html     # File Upload
        │   ├── admin.html      # Admin panel
        │   └── post.html       # IDOR vulnerability
        ├── static/             # CSS/JS files
        ├── uploads/            # Uploaded files directory
        └── instance/           # SQLite database
        """)
        
        self.doc.add_page_break()
        
        # ============================================
        # 4. FEATURES & VULNERABILITIES
        # ============================================
        self.add_heading('4. FEATURES & VULNERABILITIES', 1)
        
        self.add_heading('4.1 Security Vulnerabilities Implemented', 2)
        
        vuln_data = [
            ('#', 'Vulnerability', 'Location', 'Severity', 'Description'),
            ('1', 'SQL Injection', 'Login Page', 'Critical', 'Allows login bypass by manipulating SQL queries'),
            ('2', 'Cross-Site Scripting (XSS)', 'Search Page', 'High', 'Injects malicious scripts into web pages'),
            ('3', 'Command Injection', 'Ping Tool', 'Critical', 'Executes system commands on server'),
            ('4', 'IDOR', 'Profile/Posts', 'Medium', 'Access unauthorized user data'),
            ('5', 'Unrestricted File Upload', 'Upload Page', 'Critical', 'Uploads any file type without validation'),
            ('6', 'Weak Authentication', 'Registration', 'Medium', 'No password requirements or email validation')
        ]
        self.add_table(['#', 'Vulnerability', 'Location', 'Severity', 'Description'], vuln_data[1:])
        
        self.add_heading('4.2 Vulnerability Details', 2)
        
        # SQL Injection
        self.add_heading('SQL Injection', 3)
        self.add_paragraph("""
        What it does: Allows attackers to bypass login by manipulating SQL queries
        How to test: Use username 'admin' OR '1'='1
        Why it's bad: Gives unauthorized access to admin accounts
        Payload: admin' OR '1'='1
        """)
        
        self.add_screenshot_placeholder(2, "SQL Injection Login Form")
        
        # XSS
        self.add_heading('Cross-Site Scripting (XSS)', 3)
        self.add_paragraph("""
        What it does: Injects malicious scripts into web pages
        How to test: Search for <script>alert('XSS')</script>
        Why it's bad: Can steal cookies, redirect users, deface websites
        Payload: <script>alert('XSS Attack!')</script>
        """)
        
        self.add_screenshot_placeholder(3, "XSS Alert Popup")
        
        # Command Injection
        self.add_heading('Command Injection', 3)
        self.add_paragraph("""
        What it does: Executes system commands on the server
        How to test: Enter 'whoami' or 'ipconfig' in the ping tool
        Why it's bad: Can access files, install malware, take over server
        Payload: whoami, ipconfig, echo Hacked! & whoami
        """)
        
        self.add_screenshot_placeholder(4, "Command Injection Output")
        
        # IDOR
        self.add_heading('IDOR (Insecure Direct Object Reference)', 3)
        self.add_paragraph("""
        What it does: Allows viewing other users' data
        How to test: Go to /profile/admin or /post/1
        Why it's bad: Exposes private user information
        Payload: /profile/admin, /post/1
        """)
        
        self.add_screenshot_placeholder(5, "IDOR - Accessing Admin Profile")
        
        # File Upload
        self.add_heading('Unrestricted File Upload', 3)
        self.add_paragraph("""
        What it does: Allows uploading ANY file type
        How to test: Upload .php, .html, .exe files
        Why it's bad: Can upload malware, backdoors, scripts
        Payload: test.php, malicious.html
        """)
        
        self.add_screenshot_placeholder(6, "File Upload Interface")
        
        # Weak Auth
        self.add_heading('Weak Authentication', 3)
        self.add_paragraph("""
        What it does: No password requirements or email validation
        How to test: Register with password '123' or email 'invalid'
        Why it's bad: Weak passwords are easy to guess
        Payload: Password: 123, Email: invalidemail
        """)
        
        self.add_screenshot_placeholder(7, "Weak Password Registration")
        
        self.doc.add_page_break()
        
        # ============================================
        # 5. TESTING & DEMONSTRATION
        # ============================================
        self.add_heading('5. TESTING & DEMONSTRATION', 1)
        
        self.add_heading('5.1 Test Cases', 2)
        
        test_data = [
            ('#', 'Test Case', 'Input', 'Expected Result', 'Status'),
            ('1', 'SQL Injection', "admin' OR '1'='1", 'Successful login as admin', '✅ PASS'),
            ('2', 'XSS Attack', "<script>alert('XSS')</script>", 'Alert popup appears', '✅ PASS'),
            ('3', 'Command Injection', 'whoami', 'Show current user', '✅ PASS'),
            ('4', 'Command Injection', 'ipconfig', 'Show network info', '✅ PASS'),
            ('5', 'IDOR - Profile', '/profile/admin', 'View admin profile', '✅ PASS'),
            ('6', 'IDOR - Posts', '/post/1', 'View post #1', '✅ PASS'),
            ('7', 'File Upload', 'test.php', 'File uploaded', '✅ PASS'),
            ('8', 'Weak Password', 'Password: 123', 'Registration success', '✅ PASS'),
            ('9', 'Invalid Email', 'Email: invalid', 'Registration success', '✅ PASS'),
            ('10', 'Admin Access', '/admin', 'Admin panel shown', '✅ PASS')
        ]
        self.add_table(['#', 'Test Case', 'Input', 'Expected Result', 'Status'], test_data[1:])
        
        self.add_heading('5.2 Testing Environment', 2)
        self.add_paragraph("""
        • Operating System: Windows 10/11
        • Python Version: 3.8+
        • Browser: Google Chrome / Mozilla Firefox
        • Local Server: Flask development server
        • Database: SQLite 3
        """)
        
        self.doc.add_page_break()
        
        # ============================================
        # 6. RESULTS & OUTPUT
        # ============================================
        self.add_heading('6. RESULTS & OUTPUT', 1)
        
        self.add_heading('6.1 SQL Injection Test Results', 2)
        self.add_paragraph("""
        The SQL injection vulnerability was successfully exploited using the payload 'admin' OR '1'='1'.
        This bypassed the authentication mechanism and provided administrative access.
        """)
        self.add_screenshot_placeholder(8, "SQL Injection Successful Login")
        
        self.add_heading('6.2 XSS Test Results', 2)
        self.add_paragraph("""
        Cross-Site Scripting was confirmed by injecting the payload <script>alert('XSS')</script>
        in the search field. The script executed successfully, displaying an alert popup.
        """)
        self.add_screenshot_placeholder(9, "XSS Alert Popup")
        
        self.add_heading('6.3 Command Injection Results', 2)
        self.add_paragraph("""
        Command injection was successful using the following payloads:
        
        • whoami - Displayed the current Windows username
        • ipconfig - Displayed network configuration
        • echo Hacked! & whoami - Executed multiple commands
        """)
        self.add_screenshot_placeholder(10, "Command Injection - whoami")
        
        self.add_heading('6.4 IDOR Test Results', 2)
        self.add_paragraph("""
        IDOR vulnerabilities were confirmed by:
        
        • Accessing /profile/admin as a regular user
        • Viewing posts at /post/1, /post/2 without authorization
        """)
        self.add_screenshot_placeholder(11, "IDOR - Accessing Admin Profile")
        
        self.add_heading('6.5 File Upload Results', 2)
        self.add_paragraph("""
        The file upload feature successfully accepted all file types including:
        
        • test.php - PHP script file
        • malicious.html - HTML file with XSS payload
        • test.exe - Executable file
        """)
        self.add_screenshot_placeholder(12, "File Upload Successful")
        
        self.add_heading('6.6 Weak Authentication Results', 2)
        self.add_paragraph("""
        The authentication system was found to be weak:
        
        • Password '123' was accepted (no complexity requirements)
        • Email 'invalidemail' was accepted (no validation)
        • Multiple accounts registered with weak credentials
        """)
        self.add_screenshot_placeholder(13, "Weak Password Registration")
        
        self.doc.add_page_break()
        
        # ============================================
        # 7. LIMITATIONS
        # ============================================
        self.add_heading('7. LIMITATIONS', 1)
        
        self.add_paragraph("""
        1. Limited Vulnerability Coverage
           • Only covers 6 of OWASP Top 10 vulnerabilities
           • Does not include: CSRF, SSRF, XXE, or Deserialization flaws
        
        2. No Advanced Exploitation
           • No automatic exploitation or reporting features
           • Does not simulate real-world attack chains
        
        3. Single-User Database
           • No multi-user concurrent support
           • Data resets on restart
        
        4. No Security Controls
           • No logging or monitoring of attacks
           • No built-in defense mechanisms
        
        5. Limited Authentication
           • Uses weak hashing (SHA256 without salt)
           • No session timeout or security headers
           • No 2FA or MFA demonstration
        
        6. No API or Integration
           • Standalone application only
           • No REST API for integration with other tools
        
        7. No Mobile Support
           • Not responsive for mobile devices
           • No mobile-specific vulnerabilities
        
        8. Limited Testing Capabilities
           • Only manual testing possible
           • No automated scanning features
        """)
        
        self.doc.add_page_break()
        
        # ============================================
        # 8. FUTURE IMPROVEMENTS
        # ============================================
        self.add_heading('8. FUTURE IMPROVEMENTS', 1)
        
        self.add_paragraph("""
        1. Additional Vulnerabilities
           • Add CSRF (Cross-Site Request Forgery)
           • Add SSRF (Server-Side Request Forgery)
           • Add XXE (XML External Entity)
           • Add Deserialization vulnerabilities
        
        2. Advanced Features
           • Automatic exploitation scripts
           • Vulnerability reporting system
           • Scoring system for CTF competitions
           • Real-time attack monitoring
        
        3. Enhanced User Experience
           • Responsive mobile design
           • Interactive tutorials
           • Video demonstrations
           • Step-by-step guides
        
        4. Security Improvements
           • Add logging and monitoring
           • Implement security headers
           • Add rate limiting
           • Session management improvements
        
        5. Integration Capabilities
           • REST API for automation
           • Integration with Burp Suite and OWASP ZAP
           • Export vulnerability reports
           • Integration with SIEM systems
        """)
        
        self.doc.add_page_break()
        
        # ============================================
        # 9. CONCLUSION
        # ============================================
        self.add_heading('9. CONCLUSION', 1)
        
        self.add_paragraph("""
        Vulnerable Lab successfully demonstrates six common web application vulnerabilities in a 
        controlled, educational environment. The project serves as an effective learning tool for 
        cybersecurity students and professionals to understand:
        
        • How vulnerabilities can be exploited
        • The impact of insecure coding practices
        • The importance of security controls
        • Real-world attack vectors
        
        The application provides hands-on experience with:
        
        • SQL injection attacks
        • Cross-site scripting (XSS)
        • Command injection
        • Insecure Direct Object References (IDOR)
        • Unrestricted file uploads
        • Weak authentication mechanisms
        
        While the current implementation has limitations, it provides a solid foundation for 
        understanding web application security. The planned future improvements will enhance 
        its educational value and practical applications.
        
        This project demonstrates the critical importance of security in software development 
        and provides valuable insights for building more secure applications in the future.
        """)
        
        self.doc.add_page_break()
        
        # ============================================
        # 10. REFERENCES
        # ============================================
        self.add_heading('10. REFERENCES', 1)
        
        self.add_paragraph("""
        1. OWASP Foundation. (2021). OWASP Top Ten. https://owasp.org/www-project-top-ten/
        
        2. PortSwigger. (2023). Web Security Academy. https://portswigger.net/web-security
        
        3. Flask Documentation. (2023). Flask Web Framework. https://flask.palletsprojects.com/
        
        4. SQLAlchemy Documentation. (2023). SQLAlchemy ORM. https://www.sqlalchemy.org/
        
        5. Bootstrap Documentation. (2023). Bootstrap 5. https://getbootstrap.com/
        
        6. Python Documentation. (2023). Python 3. https://docs.python.org/
        
        7. National Institute of Standards and Technology. (2022). NIST Cybersecurity Framework.
        
        8. SANS Institute. (2022). SANS Security Training. https://www.sans.org/
        """)
        
        # ============================================
        # APPENDIX - Setup Instructions
        # ============================================
        self.doc.add_page_break()
        self.add_heading('APPENDIX A: SETUP INSTRUCTIONS', 1)
        
        self.add_heading('A.1 Prerequisites', 2)
        self.add_paragraph("""
        • Python 3.8 or higher
        • pip (Python package manager)
        • Git (optional)
        """)
        
        self.add_heading('A.2 Installation Steps', 2)
        self.add_paragraph("""
        1. Clone the repository:
           git clone <repository-url>
           cd vulnerable_lab
        
        2. Create virtual environment:
           Windows: python -m venv venv
           Windows: venv\\Scripts\\activate
           Linux/Mac: python3 -m venv venv
           Linux/Mac: source venv/bin/activate
        
        3. Install dependencies:
           pip install -r requirements.txt
        
        4. Run the application:
           python app.py
        
        5. Access the application:
           http://localhost:5000
        """)
        
        self.add_heading('A.3 Default Credentials', 2)
        cred_data = [
            ('Username', 'Password', 'Role'),
            ('admin', 'admin123', 'Administrator'),
            ('rahul', 'rahul123', 'Regular User')
        ]
        self.add_table(['Username', 'Password', 'Role'], cred_data[1:])
        
        self.add_heading('A.4 Test Payloads', 2)
        self.add_paragraph("""
        SQL Injection: admin' OR '1'='1
        XSS: <script>alert('XSS')</script>
        Command Injection: whoami, ipconfig, echo Hacked! & whoami
        IDOR: /profile/admin, /post/1
        """)

# ============================================
# MAIN EXECUTION
# ============================================
def generate_report():
    """Generate the complete report"""
    print("🚀 Generating Project Report...")
    
    # Create report generator
    generator = ReportGenerator()
    
    # Generate the report
    generator.generate_report()
    
    # Save as DOCX
    docx_filename = 'Vulnerable_Lab_Project_Report.docx'
    generator.doc.save(docx_filename)
    print(f"✅ DOCX Report saved: {docx_filename}")
    
    # Try to convert to PDF (requires python-docx2pdf)
    try:
        from docx2pdf import convert
        pdf_filename = 'Vulnerable_Lab_Project_Report.pdf'
        convert(docx_filename, pdf_filename)
        print(f"✅ PDF Report saved: {pdf_filename}")
    except ImportError:
        print("⚠️ python-docx2pdf not installed. To generate PDF:")
        print("   pip install python-docx2pdf")
        print("   Then run: python generate_report.py")
    except Exception as e:
        print(f"⚠️ Could not generate PDF: {e}")
        print("   You can still use the DOCX file.")
    
    print("\n✅ Report generation complete!")
    print(f"📄 Report saved as: {docx_filename}")
    
    # Get file size
    size = os.path.getsize(docx_filename) / 1024
    print(f"📊 File size: {size:.2f} KB")
    
    return docx_filename

if __name__ == '__main__':
    # Install required packages if not present
    try:
        import docx
    except ImportError:
        print("📦 Installing python-docx...")
        os.system('pip install python-docx')
    
    generate_report()